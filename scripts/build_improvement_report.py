"""Tạo báo cáo Word hướng dẫn cải thiện dự án PM2.5."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "Huong_dan_cai_thien_du_an_HCMC_PM25_chi_tiet.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "117A8B"
LIGHT_BLUE = "E8F1F8"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "EAF4EA"
PALE_YELLOW = "FFF4CE"
PALE_RED = "FDECEC"
GRAY = "606060"
WHITE = "FFFFFF"
BLACK = "111111"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths_dxa, strict=False)):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, *, name="Arial", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_keep(paragraph, keep_next=False, keep_lines=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        p_pr.append(OxmlElement("w:keepLines"))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=9, color=GRAY)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def add_heading(doc, text: str, level: int = 1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    set_paragraph_keep(paragraph, keep_next=True)
    return paragraph


def add_body(doc, text: str, *, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullet(doc, text: str, level: int = 0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_number(doc, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_callout(doc, title: str, text: str, fill: str = LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(f"{title}: ")
    set_run_font(title_run, bold=True, color=NAVY)
    body_run = paragraph.add_run(text)
    set_run_font(body_run, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, code: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(code.strip().splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=8.5, color="202020")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Tầng", "Trạng thái", "Kết luận", "Ưu tiên"]
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        set_cell_shading(cell, NAVY)
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, bold=True, color=WHITE, size=9.5)
    set_repeat_table_header(table.rows[0])
    for layer, status, conclusion, priority in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, [layer, status, conclusion, priority], strict=True):
            run = cell.paragraphs[0].add_run(value)
            set_run_font(run, size=9)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        fill = PALE_GREEN if status == "Đạt" else PALE_YELLOW if status == "Đạt một phần" else PALE_RED
        set_cell_shading(cells[1], fill)
    set_table_geometry(table, [1450, 1450, 4760, 1700])
    doc.add_paragraph()


def add_acceptance_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["Hạng mục", "Điều kiện nghiệm thu", "Minh chứng"], strict=True):
        set_cell_shading(cell, BLUE)
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, bold=True, color=WHITE, size=9)
    set_repeat_table_header(table.rows[0])
    for item, condition, evidence in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, [item, condition, evidence], strict=True):
            run = cell.paragraphs[0].add_run(value)
            set_run_font(run, size=8.7)
    set_table_geometry(table, [2100, 4500, 2760])
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.22

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2
    doc.styles["List Bullet"].paragraph_format.left_indent = Inches(0.375)
    doc.styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.188)
    doc.styles["List Bullet 2"].paragraph_format.left_indent = Inches(0.625)
    doc.styles["List Bullet 2"].paragraph_format.first_line_indent = Inches(-0.188)
    doc.styles["List Number"].paragraph_format.left_indent = Inches(0.375)
    doc.styles["List Number"].paragraph_format.first_line_indent = Inches(-0.188)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("HCMC PM2.5 FORECASTING  |  TECHNICAL IMPROVEMENT GUIDE")
    set_run_font(run, size=8.5, color=GRAY, bold=True)

    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("BÁO CÁO KỸ THUẬT & LỘ TRÌNH NÂNG CẤP")
    set_run_font(run, size=10, color=TEAL, bold=True)
    kicker.paragraph_format.space_after = Pt(16)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Hướng dẫn cải thiện dự án\nDự báo PM2.5 giờ tiếp theo tại TP.HCM")
    set_run_font(run, size=25, color=NAVY, bold=True)
    title.paragraph_format.space_after = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Audit theo 4 tầng: Problem → AI/ML Correctness → Software Engineering → Production & Business Value")
    set_run_font(run, size=12.5, color=DARK_BLUE)
    subtitle.paragraph_format.space_after = Pt(28)

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [7500])
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        "Mục tiêu: biến repository hiện tại thành một portfolio ML Engineering trung thực, tái lập được, không leakage và có tiêu chí triển khai rõ ràng."
    )
    set_run_font(run, size=11, color=NAVY, bold=True)

    for _ in range(5):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Phiên bản báo cáo: {date.today().strftime('%d/%m/%Y')}\nPhạm vi: repository hcmc-pm25-forecasting")
    set_run_font(run, size=9.5, color=GRAY)
    meta.add_run().add_break(WD_BREAK.PAGE)


def add_static_toc(doc: Document) -> None:
    add_heading(doc, "Mục lục", 1)
    sections = [
        "1. Tóm tắt điều hành",
        "2. Phạm vi và phương pháp đánh giá",
        "3. Tầng 1 — Problem",
        "4. Tầng 2 — AI/ML Correctness",
        "5. Tầng 3 — Software Engineering",
        "6. Tầng 4 — Production & Business Value",
        "7. Hướng dẫn sửa theo từng tệp",
        "8. Chiến lược kiểm thử và nghiệm thu",
        "9. Lộ trình triển khai ưu tiên",
        "10. Cách trình bày dự án trong CV/phỏng vấn",
        "11. Checklist phát hành GitHub",
        "12. Nguồn tham khảo",
    ]
    for item in sections:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.style = doc.styles["Normal"]
        paragraph.paragraph_format.left_indent = Inches(0.2)
        run = paragraph.add_run(item)
        set_run_font(run, color=DARK_BLUE)
    doc.add_page_break()


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)
    add_page_number(section.footer.paragraphs[0])

    add_cover(doc)
    add_static_toc(doc)

    add_heading(doc, "1. Tóm tắt điều hành", 1)
    add_body(doc, "Repository đã có nền tảng tốt cho một dự án portfolio Machine Learning Engineering: cấu trúc module rõ ràng, feature engineering theo timestamp, expanding-window backtest, hai baseline chuỗi thời gian, sklearn Pipeline, FastAPI, Streamlit, Docker và CI. Tuy nhiên, dự án chưa đủ bằng chứng để tự nhận là production-grade hoặc hệ thống cảnh báo chất lượng không khí đáng tin cậy.")
    add_callout(doc, "Kết luận chính", "Ưu tiên số một là sửa leakage một giờ tại ranh giới train/validation/test. Sau đó cần chuẩn hóa ý nghĩa ngưỡng PM2.5, yêu cầu đủ lịch sử khi inference và thay confidence heuristic bằng prediction interval đã hiệu chuẩn.", PALE_RED)
    add_status_table(doc, [
        ("Problem", "Đạt một phần", "Bài toán t+1 rõ nhưng dữ liệu mẫu chưa chứng minh use case TP.HCM thực tế.", "Cao"),
        ("AI/ML correctness", "Chưa đạt", "Feature an toàn nhưng target vượt qua ranh giới split một giờ; model thua persistence.", "Rất cao"),
        ("Software Engineering", "Đạt một phần", "Code, CI và pipeline tốt; edge-case/API tests và dependency locking chưa đủ.", "Cao"),
        ("Production/Business", "Chưa đạt", "Chưa load test, chưa có security controls và chưa chứng minh business lift.", "Trung bình"),
    ])
    add_body(doc, "Kết quả hiện tại cần được trình bày trung thực: Extra Trees có MAE xấp xỉ 2,413 trên tập test mẫu, trong khi persistence baseline đạt MAE 0,430. Quality gate đánh dấu “không đạt” là hành vi đúng; hệ thống không nên triển khai model chỉ vì đó là model tốt nhất trong nhóm ứng viên.")

    add_heading(doc, "2. Phạm vi và phương pháp đánh giá", 1)
    add_body(doc, "Báo cáo đánh giá repository theo bốn tầng liên tiếp. Một tầng phía sau chỉ có ý nghĩa khi tầng trước đã hợp lệ: bài toán đúng trước khi tối ưu mô hình; mô hình đúng trước khi đóng gói phần mềm; phần mềm ổn định trước khi tuyên bố giá trị production.")
    flow = doc.add_table(rows=4, cols=1)
    set_table_geometry(flow, [7000])
    flow.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, (text, fill) in enumerate([
        ("1  PROBLEM — mục tiêu, input/output, use case", "DCE6F1"),
        ("2  AI/ML CORRECTNESS — leakage, split, metric, baseline", "D9EAD3"),
        ("3  SOFTWARE ENGINEERING — code, test, API, CI, tái lập", "FFF2CC"),
        ("4  PRODUCTION / BUSINESS — tải, bảo mật, giám sát, giá trị", "F4CCCC"),
    ]):
        cell = flow.cell(index, 0)
        set_cell_shading(cell, fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, bold=True, color=NAVY)
    doc.add_paragraph()

    add_heading(doc, "3. Tầng 1 — Problem", 1)
    add_heading(doc, "3.1. Định nghĩa bài toán chuẩn", 2)
    add_callout(doc, "Problem statement đề xuất", "Tại thời điểm t, sử dụng toàn bộ quan trắc đã biết đến hết t của một trạm để dự báo nồng độ PM2.5 tại t+1 giờ. Không sử dụng bất kỳ feature hoặc target nào có timestamp sau thời điểm dự báo.")
    add_heading(doc, "3.2. Input hợp lệ", 2)
    for text in [
        "Request chỉ chứa một trạm; station không rỗng và có độ dài hợp lý.",
        "Timestamp hợp lệ, tăng dần, không trùng và có timezone hoặc quy ước timezone rõ ràng.",
        "PM2.5 tại thời điểm hiện tại phải tồn tại, hữu hạn và không âm.",
        "Nếu dùng lag/rolling 24 giờ, yêu cầu ít nhất 25 mốc liên tục; không nên nhận hai dòng rồi impute gần như toàn bộ lịch sử.",
        "O3, SO2 và biến ngoại sinh có thể thiếu nhưng response phải phản ánh chất lượng dữ liệu.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "3.3. Output đề xuất", 2)
    add_code(doc, '''{
  "station": "Trạm A",
  "forecast_origin": "2024-01-01T10:00:00+07:00",
  "forecast_for": "2024-01-01T11:00:00+07:00",
  "current_pm25": 32.1,
  "predicted_pm25": 34.7,
  "level": "Trung bình",
  "prediction_interval": {"lower": 29.2, "upper": 40.1, "coverage": 0.90},
  "data_quality": {"missing_hours": 0, "imputed_features": 2},
  "model_version": "2026-09-01-001",
  "updated_at": "2026-09-01T08:00:00Z"
}''')
    add_body(doc, "`forecast_origin` là mốc cuối cùng đã quan sát; `forecast_for` là thời điểm mục tiêu. `updated_at` chỉ là lúc API tạo response. Việc tách ba timestamp giúp demo không gây hiểu nhầm.")
    add_heading(doc, "3.4. Phân biệt hồi quy và phân loại mức", 2)
    add_body(doc, "Mục tiêu học máy chính là hồi quy nồng độ PM2.5. Nhãn mức chỉ là lớp diễn giải sau dự báo. Metric hồi quy quyết định chất lượng dự báo; metric phân lớp chỉ hợp lệ khi ngưỡng được định nghĩa, trích nguồn và áp dụng đúng thời lượng trung bình.")

    add_heading(doc, "4. Tầng 2 — AI/ML Correctness", 1)
    add_heading(doc, "4.1. Lỗi leakage tại biên target", 2)
    add_body(doc, "Pipeline hiện tạo target_next_hour trước rồi chia dữ liệu dựa trên timestamp của feature. Vì vậy, dòng train cuối tại 10:00 có thể dùng PM2.5 lúc 11:00 làm target, trong khi test bắt đầu tại 11:00. Đây là leakage nhãn một giờ tại biên train/test. Mỗi expanding-window fold có rủi ro tương tự.")
    add_code(doc, '''# Bổ sung thời gian thật của target
result["target_timestamp"] = (
    result[timestamp_column] + pd.Timedelta(hours=1)
)

# Train chỉ được dùng target đã xảy ra trước test
train_mask = frame["target_timestamp"] < test_start
test_mask = frame[timestamp_column] >= test_start''')
    add_callout(doc, "Điều kiện bất biến", "max(train.target_timestamp) phải nhỏ hơn min(validation.timestamp) và min(test.timestamp). Kiểm tra chỉ dựa trên feature timestamp là chưa đủ.", PALE_RED)
    add_heading(doc, "4.2. Feature engineering", 2)
    add_body(doc, "Cách tra lag theo khóa (station, timestamp) là đúng và tốt hơn shift theo vị trí dòng. Nếu thiếu đúng một giờ, lag phải là NaN thay vì lấy nhầm dòng gần nhất. Rolling sử dụng closed='left' nên loại trừ quan trắc hiện tại; model vẫn có thể dùng PM2.5 hiện tại như một feature riêng để dự báo t+1.")
    add_heading(doc, "4.3. Ngưỡng PM2.5 và nguy cơ diễn giải sai", 2)
    add_body(doc, "Ngưỡng 12,0 và 35,5 gần với breakpoint PM2.5 AQI của Hoa Kỳ trước cập nhật năm 2024. EPA đã thay breakpoint Good từ 12,0 xuống 9,0 µg/m³. Quan trọng hơn, breakpoint AQI PM2.5 được áp dụng cho giá trị trung bình 24 giờ, không nên gắn trực tiếp vào một dự báo nồng độ từng giờ rồi gọi là AQI chính thức.")
    add_callout(doc, "Khuyến nghị", "Với portfolio, đổi nhãn thành Thấp/Trung bình/Cao và ghi rõ đây là phân nhóm thử nghiệm. Nếu muốn hiển thị AQI, phải tính rolling 24 giờ và áp dụng công thức/breakpoint của tiêu chuẩn được chọn.", PALE_YELLOW)
    add_heading(doc, "4.4. Split và validation", 2)
    for text in [
        "Giữ test cuối bất biến và chỉ sử dụng một lần sau khi chọn model.",
        "Chọn model/hyperparameter bằng expanding-window backtest chỉ trên phần train.",
        "Giữ tất cả trạm tại cùng timestamp trong cùng partition.",
        "Tạo optional gap một giờ giữa train và validation/test để tăng độ thận trọng.",
        "Với dữ liệu thật, ưu tiên mốc ngày cố định thay vì test_fraction để báo cáo dễ tái lập.",
    ]:
        add_bullet(doc, text)
    add_code(doc, '''split:
  train_end: "2023-12-31 23:00:00"
  validation_end: "2024-03-31 23:00:00"
  test_end: "2024-06-30 23:00:00"
  gap_hours: 1
  backtest_folds: 5''')
    add_heading(doc, "4.5. Metric", 2)
    add_acceptance_table(doc, [
        ("MAE", "Metric chính, dễ diễn giải theo µg/m³", "Toàn test, từng trạm, từng mùa"),
        ("RMSE", "Phạt mạnh sai số lớn", "Báo cùng MAE, không dùng đơn độc"),
        ("Bias", "Mean(y_pred − y_true) gần 0", "Phát hiện model dự báo thấp/cao có hệ thống"),
        ("P90 absolute error", "90% sai số tuyệt đối không vượt ngưỡng báo cáo", "Đo tail error dễ hiểu"),
        ("Macro-F1/QWK", "Chỉ dùng khi ngưỡng lớp hợp lệ", "Confusion matrix kèm support"),
        ("High-PM2.5 recall", "Không được che giấu tỷ lệ bỏ sót", "Báo recall và số FN tuyệt đối"),
        ("Rolling MAE", "Báo mean ± std", "Cho thấy biến động theo giai đoạn"),
    ])
    add_heading(doc, "4.6. Baseline và quality gate", 2)
    add_body(doc, "Persistence và seasonal naive 24 giờ là baseline phù hợp. Quality gate cần được mở rộng để model không chỉ thắng về MAE mà còn đủ ổn định và không bỏ sót quá nhiều sự kiện cao.")
    add_code(doc, '''quality_gate:
  minimum_mae_improvement: 0.05
  minimum_high_pm25_recall: 0.75
  maximum_rolling_mae_std: 1.00''')
    add_body(doc, "Nếu không model nào vượt persistence, champion production nên là persistence. Đây là kết luận khoa học hợp lệ, không phải thất bại của dự án.")
    add_heading(doc, "4.7. Thay confidence heuristic", 2)
    add_body(doc, "Công thức 1 − std(các cây)/prediction không phải xác suất model đúng và không có coverage guarantee. Nếu chưa có hiệu chuẩn, nên xóa trường confidence hoặc đổi thành tree_agreement_score. Giải pháp tốt hơn là split conformal interval dựa trên residual validation.")
    add_code(doc, '''residuals = np.abs(y_validation - validation_prediction)
q90 = np.quantile(residuals, 0.90)

lower = max(0.0, prediction - q90)
upper = prediction + q90''')

    add_heading(doc, "5. Tầng 3 — Software Engineering", 1)
    add_heading(doc, "5.1. Validation inference", 2)
    add_code(doc, '''required_hours = max(
    max(config["features"]["lags"]),
    max(config["features"]["rolling_windows"]),
)

if len(observations) < required_hours + 1:
    raise InsufficientHistoryError(
        f"Cần tối thiểu {required_hours + 1} quan trắc."
    )

gaps = observations.sort_values(timestamp_column)[timestamp_column].diff().dropna()
if (gaps != pd.Timedelta(hours=1)).any():
    raise IrregularHistoryError("Chuỗi quan trắc không liên tục theo giờ.")''')
    add_heading(doc, "5.2. Pydantic schema", 2)
    add_code(doc, '''class Observation(BaseModel):
    timestamp: datetime
    station: str = Field(min_length=1, max_length=100)
    PM25: float = Field(alias="PM2.5", ge=0, le=1000)
    temperature: float | None = Field(default=None, ge=-20, le=60)
    humidity: float | None = Field(default=None, ge=0, le=100)

class PredictionRequest(BaseModel):
    observations: list[Observation] = Field(min_length=25, max_length=168)''')
    add_heading(doc, "5.3. Chuẩn hóa lỗi API", 2)
    add_acceptance_table(doc, [
        ("400", "Input có logic không hợp lệ", "INSUFFICIENT_HISTORY, MULTIPLE_STATIONS"),
        ("422", "Pydantic schema thất bại", "Giá trị âm, timestamp sai, thiếu trường"),
        ("503", "Model hoặc artifact chưa sẵn sàng", "MODEL_NOT_READY"),
        ("500", "Lỗi nội bộ không dự kiến", "Không trả stack trace/đường dẫn hệ thống"),
    ])
    add_heading(doc, "5.4. Health check", 2)
    add_body(doc, "Health endpoint hiện chỉ chứng minh tiến trình FastAPI còn sống. Cần tách liveness và readiness: liveness kiểm tra process; readiness xác nhận config, model artifact và predictor đã nạp thành công.")
    add_heading(doc, "5.5. Dependency và tái lập", 2)
    for text in [
        "Giữ requirements.in để con người chỉnh dependency cấp cao.",
        "Sinh requirements.lock hoặc uv.lock với phiên bản chính xác.",
        "Docker và CI phải cài từ lock file.",
        "Lưu Python, NumPy, pandas và scikit-learn version trong metadata.",
        "Giữ notebook Colab dùng cùng src.train và snapshot/hash dữ liệu.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "5.6. Tách training khỏi Docker serving", 2)
    add_body(doc, "Không nên train model trong Docker build. Build image phải nhanh và ổn định; training là một job riêng sinh artifact. Serving image chỉ nạp artifact đã vượt quality gate và được xác minh hash.")
    add_code(doc, '''Training job
    │
    ├── model.joblib
    ├── metadata.json
    └── evaluation.json
            │
            ▼
Serving image → FastAPI → Dashboard''')
    add_heading(doc, "5.7. Metadata artifact", 2)
    add_code(doc, '''{
  "model_version": "2026-09-01-001",
  "git_commit": "<commit-sha>",
  "data_sha256": "<hash>",
  "config_sha256": "<hash>",
  "python_version": "3.11.x",
  "scikit_learn_version": "1.5.2",
  "train_period": ["...", "..."],
  "test_period": ["...", "..."],
  "quality_gate": {"status": "đạt|không đạt"}
}''')

    add_heading(doc, "6. Tầng 4 — Production & Business Value", 1)
    add_heading(doc, "6.1. Load test 100 users", 2)
    add_body(doc, "Không thể khẳng định hệ thống chịu được 100 người dùng nếu chưa đo. Dùng Locust hoặc k6 với payload đại diện và chạy trên cùng cấu hình container dự định demo.")
    add_acceptance_table(doc, [
        ("Concurrency", "100 virtual users", "Kịch bản ramp-up và steady-state"),
        ("Error rate", "< 1%", "Loại trừ lỗi client cố ý"),
        ("p95 latency", "< 500 ms", "Endpoint /predict"),
        ("p99 latency", "< 1.000 ms", "Endpoint /predict"),
        ("Health latency", "< 100 ms", "Readiness không thực hiện inference"),
        ("Stability", "Không tăng RAM không kiểm soát", "Theo dõi trong 10–15 phút"),
    ])
    add_heading(doc, "6.2. Security tối thiểu", 2)
    for text in [
        "Giới hạn kích thước payload và số observation.",
        "Thêm rate limiting ở reverse proxy/API gateway.",
        "Cấu hình CORS theo domain dashboard, không dùng wildcard khi public.",
        "Không trả exception nội bộ cho client; dùng request ID để tra log.",
        "Chỉ nạp joblib artifact từ nguồn tin cậy và xác minh hash.",
        "Chạy container non-root; không commit .env hoặc secrets.toml.",
        "Chạy pip-audit và Bandit trong CI.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "6.3. Monitoring", 2)
    add_body(doc, "Monitoring nên tập trung vào hành vi hệ thống và chất lượng dữ liệu/model: request count, error rate, latency, tỷ lệ feature thiếu, phân phối PM2.5, drift theo trạm, độ rộng interval và MAE sau khi target thật xuất hiện. Portfolio không bắt buộc có Prometheus/Grafana; structured JSON logs và một báo cáo định kỳ đã đủ hợp lý.")
    add_heading(doc, "6.4. Dashboard", 2)
    for text in [
        "Hiển thị forecast origin và forecast for thay vì chỉ updated_at.",
        "Hiển thị prediction interval; không gọi heuristic là độ tin cậy.",
        "Hiển thị cảnh báo lịch sử thiếu hoặc feature đã impute.",
        "Hiển thị model version và quality gate.",
        "Phân biệt lỗi 422, 503 và lỗi kết nối API.",
        "Giữ disclaimer: không dùng thay cho cảnh báo sức khỏe chính thức.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "7. Hướng dẫn sửa theo từng tệp", 1)
    add_acceptance_table(doc, [
        ("src/features.py", "Tạo target_timestamp; giữ lag theo timestamp; ghi rõ seasonal t−23 cho dự báo t+1", "Test exact-hour lag và target timestamp"),
        ("src/train.py", "Split/backtest theo target_timestamp; mở rộng quality gate", "Không có target train vượt biên"),
        ("src/evaluate.py", "Thêm bias, P90 error; đổi nhãn; xử lý support=0", "JSON strict, metric từng trạm"),
        ("src/predict.py", "Kiểm tra lịch sử; interval; model/config compatibility", "Edge-case inference tests"),
        ("app/api.py", "Schema bounds; error model; liveness/readiness; payload limit", "TestClient cho 400/422/503"),
        ("app/dashboard.py", "Phân biệt lỗi; interval; data quality; model version", "Manual smoke test/Playwright nếu cần"),
        ("configs/config.yaml", "Ngưỡng có source; split date/gap; quality gate", "validate_config chặn cấu hình sai"),
        ("Dockerfile", "Không train trong build; dùng lock file và non-root", "Docker build + health smoke test"),
        ("README.md", "Giảm overclaim; bảng baseline; limitation; Colab/Docker", "Fresh-clone instructions"),
        ("tests/", "Bổ sung leakage target, API/inference và reproducibility", "CI xanh trên Python 3.11"),
    ])

    add_heading(doc, "8. Chiến lược kiểm thử và nghiệm thu", 1)
    add_heading(doc, "8.1. Test ML bắt buộc", 2)
    tests = [
        "Train target không vượt thời điểm bắt đầu validation/test.",
        "Lag một giờ trả NaN khi thiếu đúng timestamp, không lấy dòng trước đó.",
        "Rolling window loại trừ quan trắc hiện tại.",
        "Mọi trạm cùng timestamp nằm chung partition.",
        "Chạy train hai lần cùng seed cho cùng champion và metric trong tolerance.",
        "Model không được đánh dấu đạt khi thua persistence.",
    ]
    for item in tests:
        add_bullet(doc, item)
    add_heading(doc, "8.2. Test inference/API bắt buộc", 2)
    for item in [
        "Thiếu O3/SO2 vẫn chạy qua toàn pipeline, không chỉ test SimpleImputer riêng.",
        "Từ chối nhiều trạm, timestamp trùng, timestamp sai và lịch sử không đủ.",
        "Từ chối PM2.5 âm, NaN/Infinity và payload vượt giới hạn.",
        "Station chưa gặp khi train vẫn dự báo được nhờ OneHotEncoder(handle_unknown='ignore').",
        "Thiếu/hỏng artifact trả 503 thay vì stack trace.",
        "Response JSON đúng schema và chứa forecast_origin/forecast_for/model_version.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "8.3. Lệnh nghiệm thu", 2)
    add_code(doc, '''python -m ruff check src app tests
python -m pytest -q
python -m src.train --config configs/config.yaml --dry-run
docker build -t hcmc-pm25 .
docker compose up --build
# Sau đó chạy smoke test và load test trên /health, /ready, /predict''')

    add_heading(doc, "9. Lộ trình triển khai ưu tiên", 1)
    add_heading(doc, "Giai đoạn 1 — Bắt buộc trước khi đưa lên CV", 2)
    for text in [
        "Sửa leakage target tại mọi ranh giới split.",
        "Thêm test bất biến theo target_timestamp.",
        "Đổi nhãn hoặc giải thích nguồn/ngữ cảnh ngưỡng PM2.5.",
        "Xóa hoặc đổi confidence heuristic.",
        "Yêu cầu đủ lịch sử inference và chuẩn hóa lỗi API.",
        "Chạy lại training, notebook và cập nhật toàn bộ metric.",
    ]:
        add_number(doc, text)
    add_heading(doc, "Giai đoạn 2 — Hoàn thiện portfolio", 2)
    for text in [
        "Thêm conformal prediction interval.",
        "Khóa dependency và tách training khỏi Docker serving.",
        "Thêm artifact version/hash và readiness check.",
        "Bổ sung test toàn pipeline, Docker smoke test và ảnh dashboard.",
        "Viết bảng kết quả theo model, trạm, lớp và giai đoạn.",
    ]:
        add_number(doc, text)
    add_heading(doc, "Giai đoạn 3 — Production demo", 2)
    for text in [
        "Load test 100 users và ghi p95/p99/error rate.",
        "Thêm rate limiting, CORS, payload limit và security scan.",
        "Structured logging, drift/data-quality monitoring.",
        "Thay dữ liệu mẫu bằng dữ liệu thật có nguồn và giấy phép.",
        "Đánh giá nhiều tháng, mùa, trạm và sự kiện PM2.5 cao.",
    ]:
        add_number(doc, text)

    add_heading(doc, "10. Cách trình bày dự án trong CV/phỏng vấn", 1)
    add_callout(doc, "Câu mô tả nên dùng", "Xây dựng pipeline dự báo PM2.5 giờ tiếp theo theo trạm với feature engineering theo timestamp, expanding-window validation, baseline quality gate, FastAPI/Streamlit, Docker và CI; chủ động ngăn triển khai khi model không vượt persistence baseline.", PALE_GREEN)
    add_body(doc, "Điểm mạnh khi phỏng vấn không phải là cố chứng minh Extra Trees tốt. Điểm mạnh là bạn phát hiện model thua baseline, xác định leakage ở biên target, thiết kế quality gate và biết giới hạn của dữ liệu mẫu.")
    add_heading(doc, "Các tuyên bố nên tránh", 2)
    for text in [
        "Production-grade nếu chưa có load test, monitoring và security controls.",
        "Chống leakage tuyệt đối trước khi sửa ranh giới target.",
        "Hệ thống cảnh báo chính xác khi recall sự kiện cao chưa đạt.",
        "Độ tin cậy 90% nếu chưa có interval calibration/coverage test.",
        "AQI chính thức nếu đang phân lớp nồng độ từng giờ bằng breakpoint 24 giờ.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "11. Checklist phát hành GitHub", 1)
    checklist = [
        ("☐", "Problem statement và horizon t+1 rõ ràng"),
        ("☐", "Input/output schema và timezone rõ ràng"),
        ("☐", "Không có feature hoặc target leakage"),
        ("☐", "Train/validation/test chia theo target timestamp"),
        ("☐", "Model được so với persistence và seasonal baseline"),
        ("☐", "Metric gồm MAE, RMSE, bias, tail/high-event metrics"),
        ("☐", "Ngưỡng phân loại có nguồn và đúng ngữ cảnh"),
        ("☐", "Inference yêu cầu đủ lịch sử và xử lý missing"),
        ("☐", "API trả lỗi 4xx/5xx đúng, không lộ exception nội bộ"),
        ("☐", "Không có secret hoặc path máy cá nhân"),
        ("☐", "Fresh clone chạy test, dry-run, API và dashboard"),
        ("☐", "Docker build không tự train model"),
        ("☐", "Notebook Colab dùng cùng source code và khớp snapshot"),
        ("☐", "README có kết quả trung thực và limitation"),
        ("☐", "Load test/security scan nếu tuyên bố production readiness"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_shading(table.cell(0, 0), NAVY)
    set_cell_shading(table.cell(0, 1), NAVY)
    for cell, text in zip(table.rows[0].cells, ["Trạng thái", "Điều kiện"], strict=True):
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, bold=True, color=WHITE)
    for mark, text in checklist:
        cells = table.add_row().cells
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(mark), size=12, color=BLUE)
        set_run_font(cells[1].paragraphs[0].add_run(text), size=9.2)
    set_table_geometry(table, [1200, 8160])

    add_heading(doc, "12. Nguồn tham khảo", 1)
    sources = [
        "US EPA — AQI Breakpoints: https://aqs.epa.gov/aqsweb/documents/codetables/aqi_breakpoints.html",
        "US EPA — Final Updates to the AQI for Particulate Matter (2024): https://www.epa.gov/system/files/documents/2024-02/pm-naaqs-air-quality-index-fact-sheet.pdf",
        "FastAPI Documentation — Deployment and Docker: https://fastapi.tiangolo.com/deployment/docker/",
        "scikit-learn Documentation — Pipeline, TimeSeriesSplit and metrics: https://scikit-learn.org/stable/",
        "Repository source reviewed: src/, app/, tests/, configs/, Dockerfile, docker-compose.yml, README.md and GitHub Actions workflow.",
    ]
    for source in sources:
        add_bullet(doc, source)
    add_callout(doc, "Lưu ý sử dụng", "Báo cáo này là hướng dẫn kỹ thuật. Các ngưỡng chất lượng không khí và nội dung liên quan sức khỏe phải được xác minh theo tiêu chuẩn chính thức mà dự án lựa chọn trước khi công bố.", PALE_YELLOW)

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    return doc


if __name__ == "__main__":
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.core_properties.title = "Hướng dẫn cải thiện dự án HCMC PM2.5 Forecasting"
    document.core_properties.subject = "Audit 4 tầng và lộ trình nâng cấp kỹ thuật"
    document.core_properties.author = "HCMC PM2.5 Forecasting Project"
    document.core_properties.keywords = "PM2.5, time series, leakage, FastAPI, MLOps, portfolio"
    document.save(OUTPUT)
    print(OUTPUT)
